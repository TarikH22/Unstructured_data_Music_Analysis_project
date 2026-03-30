import os
import random
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook

DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', 'data', 'raw')
os.makedirs(DATA_DIR, exist_ok=True)

ARTISTS = [
    {"name": "The Weeknd", "listeners": 1500000, "playcount": 50000000, "bio": "A Canadian singer, songwriter, and record producer known for his sonic versatility and dark lyricism."},
    {"name": "Taylor Swift", "listeners": 2000000, "playcount": 60000000, "bio": "An American singer-songwriter whose discography spans genres and whose narrative songwriting is often inspired by her personal life."},
    {"name": "Daft Punk", "listeners": 1800000, "playcount": 45000000, "bio": "A French electronic music duo formed in 1993 by Guy-Manuel de Homem-Christo and Thomas Bangalter."},
    {"name": "Radiohead", "listeners": 1200000, "playcount": 30000000, "bio": "An English rock band formed in Abingdon, Oxfordshire, in 1985."},
]

def generate_pdf():
    pdf_path = os.path.join(DATA_DIR, "artists_report.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, "Monthly Artists Report")
    
    c.setFont("Helvetica", 12)
    y_position = height - 100
    
    for artist in ARTISTS:
        c.drawString(50, y_position, f"Artist: {artist['name']}")
        c.drawString(50, y_position - 20, f"Listeners: {artist['listeners']}")
        c.drawString(50, y_position - 40, f"Playcount: {artist['playcount']}")
        c.drawString(50, y_position - 60, f"Bio: {artist['bio'][:60]}...")
        y_position -= 100
        
        if y_position < 100:
            c.showPage()
            y_position = height - 50
            
    c.save()
    print(f"Generated {pdf_path}")

def generate_docx():
    docx_path = os.path.join(DATA_DIR, "artists_summary.docx")
    doc = Document()
    doc.add_heading('Artists Summary', 0)
    
    for artist in ARTISTS:
        doc.add_heading(artist['name'], level=1)
        doc.add_paragraph(f"Listeners: {artist['listeners']}")
        doc.add_paragraph(f"Playcount: {artist['playcount']}")
        doc.add_paragraph(artist['bio'])
        
    doc.save(docx_path)
    print(f"Generated {docx_path}")

def generate_xlsx():
    xlsx_path = os.path.join(DATA_DIR, "artists_data.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Artists"
    
    ws.append(["Name", "Listeners", "Playcount", "Biography"])
    
    for artist in ARTISTS:
        ws.append([artist['name'], artist['listeners'], artist['playcount'], artist['bio']])
        
    wb.save(xlsx_path)
    print(f"Generated {xlsx_path}")

if __name__ == "__main__":
    generate_pdf()
    generate_docx()
    generate_xlsx()
