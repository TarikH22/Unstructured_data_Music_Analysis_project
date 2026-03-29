import os
from datetime import datetime

import chardet
import openpyxl
import pdfplumber
from docx import Document
from openpyxl.worksheet.table import Table, TableStyleInfo

from utils.logger import logger


ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.."))
RAW_DIR = os.path.join(ROOT_DIR, "data", "raw")
DOCS_DIR = os.path.join(ROOT_DIR, "data", "processed", "documents")


def _metadata(file_path, doc_type, source, page_number=None):
    return {
        "file_name": os.path.basename(file_path),
        "document_type": doc_type,
        "source": source,
        "page_number": page_number,
        "extraction_timestamp": datetime.utcnow().isoformat() + "Z",
    }


def ensure_docx_exists():
    os.makedirs(DOCS_DIR, exist_ok=True)
    path = os.path.join(DOCS_DIR, "LastFM_Normal.docx")
    if os.path.exists(path):
        return path

    logger.info("Generating LastFM_Normal.docx")
    song_doc = Document()
    song_doc.add_heading("Last.fm Document (Word)", level=1)
    song_doc.add_paragraph(
        "This document is generated for document extraction tests in the Last.fm pipeline."
    )

    bordered = song_doc.add_table(rows=1, cols=3)
    bordered.style = "Table Grid"
    bordered.rows[0].cells[0].text = "Artist"
    bordered.rows[0].cells[1].text = "Track"
    bordered.rows[0].cells[2].text = "Playcount"
    row = bordered.add_row().cells
    row[0].text = "Coldplay"
    row[1].text = "The Scientist"
    row[2].text = "75000000"

    borderless = song_doc.add_table(rows=2, cols=2)
    borderless.rows[0].cells[0].text = "Tag"
    borderless.rows[0].cells[1].text = "alternative rock"
    borderless.rows[1].cells[0].text = "Mood"
    borderless.rows[1].cells[1].text = "melancholic"

    image_candidates = [
        os.path.join(DOCS_DIR, "images", "local_cover.png"),
        os.path.join(DOCS_DIR, "images", "track_stats.png"),
    ]
    for image_path in image_candidates:
        if os.path.exists(image_path):
            song_doc.add_picture(image_path)

    song_doc.save(path)
    return path


def ensure_excel_exists():
    os.makedirs(DOCS_DIR, exist_ok=True)
    path = os.path.join(DOCS_DIR, "LastFM_Data.xlsx")
    if os.path.exists(path):
        return path

    logger.info("Generating LastFM_Data.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TrackStats"
    ws.append(["Artist", "Track", "Listeners", "Playcount", "EngagementScore"])
    ws.append(["Coldplay", "The Scientist", 2800000, 75000000, "=D2/C2"])
    ws.append(["Radiohead", "Creep", 2100000, 51000000, "=D3/C3"])
    ws.append(["The Weeknd", "Blinding Lights", 3400000, 98000000, "=D4/C4"])

    table = Table(displayName="TrackTable", ref="A1:E4")
    style = TableStyleInfo(
        name="TableStyleMedium9",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    table.tableStyleInfo = style
    ws.add_table(table)

    wb.save(path)
    return path


def detect_and_read_text(file_path):
    with open(file_path, "rb") as f:
        raw = f.read()
    encoding_info = chardet.detect(raw)
    encoding = encoding_info.get("encoding") or "utf-8"
    text = raw.decode(encoding, errors="replace")
    return text, encoding


def extract_pdf(pdf_path, source_label, two_column=False):
    records = []
    logger.info(f"Extracting PDF: {pdf_path}")
    with pdfplumber.open(pdf_path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            if two_column:
                half = page.width / 2
                left_text = (page.crop((0, 0, half, page.height)).extract_text() or "").strip()
                right_text = (page.crop((half, 0, page.width, page.height)).extract_text() or "").strip()
                text = (left_text + "\n" + right_text).strip()
            else:
                text = (page.extract_text() or "").strip()

            tables = page.extract_tables() or []
            records.append(
                {
                    "content": text,
                    "tables": tables,
                    "metadata": _metadata(pdf_path, "pdf", source_label, page_number=index),
                }
            )
    return records


def extract_docx(docx_path, source_label):
    logger.info(f"Extracting DOCX: {docx_path}")
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    return [
        {
            "content": "\n".join(paragraphs),
            "paragraphs": paragraphs,
            "tables": tables,
            "metadata": _metadata(docx_path, "docx", source_label),
        }
    ]


def extract_xlsx(xlsx_path, source_label):
    logger.info(f"Extracting XLSX: {xlsx_path}")
    wb = openpyxl.load_workbook(xlsx_path, data_only=False)
    records = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            rows.append(["" if cell is None else str(cell) for cell in row])

        records.append(
            {
                "content": "\n".join([", ".join(r) for r in rows]),
                "rows": rows,
                "metadata": {
                    **_metadata(xlsx_path, "xlsx", source_label),
                    "worksheet": sheet.title,
                },
            }
        )
    return records


def process_documents():
    normal_pdf = os.path.join(DOCS_DIR, "LastFM_Normal.pdf")
    two_col_pdf = os.path.join(DOCS_DIR, "LastFM_TwoColumn.pdf")
    docx_file = ensure_docx_exists()
    xlsx_file = ensure_excel_exists()

    if not os.path.exists(normal_pdf) or not os.path.exists(two_col_pdf):
        raise FileNotFoundError(
            "Missing PDFs in data/processed/documents. Run generate_docs.py first."
        )

    extracted = []
    extracted.extend(extract_pdf(normal_pdf, "normal-pdf", two_column=False))
    extracted.extend(extract_pdf(two_col_pdf, "two-column-pdf", two_column=True))
    extracted.extend(extract_docx(docx_file, "word-document"))
    extracted.extend(extract_xlsx(xlsx_file, "excel-document"))

    review_file = os.path.join(RAW_DIR, "reviews", "coldplay_my_universe_review.txt")
    if os.path.exists(review_file):
        review_text, encoding = detect_and_read_text(review_file)
        extracted.append(
            {
                "content": review_text,
                "metadata": {
                    **_metadata(review_file, "txt", "encoding-check"),
                    "detected_encoding": encoding,
                },
            }
        )

    logger.info(f"Extracted {len(extracted)} document records")
    return extracted
